"""
Regenera los documentos de planificación y funcionamiento de GazeDash
al estado actual v0.3 usando python-docx.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

GREEN = RGBColor(0x10, 0xB9, 0x81)

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    return p

def body(doc, text):
    return doc.add_paragraph(text)

def table_row(table, *cells):
    row = table.add_row()
    for i, c in enumerate(cells):
        row.cells[i].text = c
    return row

# ─────────────────────────────────────────────────────────────────────────────
# DOC 1: Cómo Funciona 0.3
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

h1(doc, "GazeDash — Cómo funciona (v0.3)")
body(doc, f"Última actualización: {datetime.date.today().strftime('%d/%m/%Y')}")
body(doc,
    "GazeDash es una aplicación de accesibilidad para Windows que permite controlar la "
    "computadora mediante gestos faciales, movimiento de nariz/cabeza como joystick y "
    "comandos de voz, usando únicamente una cámara web estándar.")

h2(doc, "Flujo general")
steps = [
    "La cámara captura video en un hilo de fondo (CameraStream).",
    "MediaPipeFaceDetector extrae landmarks, blendshapes, métricas geométricas (ratios) "
     "y la posición de la nariz (nose_tip).",
    "FacialGestureDetector convierte las métricas en gestos discretos usando calibración "
     "neutral adaptativa (primeros ~5 s) y debouncing por frames consecutivos.",
    "Los gestos activos se mapean a acciones del perfil activo (hotkeys, clicks) "
     "mediante gesture_mapper y HotkeyExecutor (pydirectinput).",
    "MouseController mueve el cursor según la posición relativa de la nariz respecto "
     "a un centro recalibrable. Aplica zona muerta, velocidad máxima y suavizado EMA.",
    "BlinkClickController detecta guiños unilaterales (un ojo cerrado, el otro abierto) "
     "y ejecuta click izquierdo (guiño izq.) o derecho (guiño der.).",
    "VoiceCommandController escucha audio en loop, clasifica módulo y comando con modelos "
     "ML (.pkl) y ejecuta la acción correspondiente.",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}", style="List Number")

h2(doc, "Detección de gestos faciales")
body(doc,
    "El sistema usa MediaPipe Face Landmarker con blendshapes (ARKit style). "
    "Los gestos disponibles son:")
gestos = [
    "mouth_open — boca abierta (jaw open + ratio geométrico)",
    "mouth_pucker — fruncir labios",
    "mouth_o (funnel) — boca en 'O'",
    "smile — sonrisa bilateral (conjunción de smile_left y smile_right)",
    "smile_left, smile_right — sonrisa asimétrica",
    "brow_raise — subir cejas",
    "brow_frown — fruncir cejas",
    "eye_blink — guiño (solo activa si es unilateral)",
    "eye_wide — ojos muy abiertos",
    "nose_sneer — arrugar la nariz",
]
for g in gestos:
    doc.add_paragraph(g, style="List Bullet")

body(doc,
    "Cada gesto tiene: escala de umbral, delta mínimo absoluto (evita ruido de baseline) "
    "y contador de frames de debounce (N frames consecutivos antes de reportar). "
    "La base neutral se adapta lentamente (adaptation_alpha=0.01) cuando ningún gesto está activo.")

h2(doc, "Mouse por nariz")
body(doc,
    "MouseController calcula dx/dy entre la posición actual de la nariz y el centro "
    "de referencia. Aplica zona muerta, velocidad máxima y suavizado (EMA). "
    "El centro se recalibra desde el botón 'Recalibrar' en la ventana principal, "
    "en la vista Calibración, o en el mini overlay. El toggle de habilitación está en "
    "Configuración → tab Mouse (CTkSwitch que guarda y aplica inmediatamente).")

h2(doc, "Clic por guiño")
body(doc,
    "BlinkClickController detecta guiños unilaterales: un ojo se cierra mientras el otro "
    "permanece abierto. Guiño izquierdo → click izquierdo; guiño derecho → click derecho. "
    "Incluye cooldown para evitar disparo doble.")

h2(doc, "Módulo de voz")
body(doc, "VoiceCommandController implementa una máquina de estados:")
estados = [
    "DISABLED — voz desactivada (toggle en Configuración → Control de voz).",
    "WAITING_MODULE — escucha continua para identificar el módulo hablado.",
    "MODULE_ACTIVE — módulo reconocido; escucha comandos de ese módulo.",
    "LISTENING_COMMAND — grabando audio para clasificar el comando.",
    "CONFIRM_REQUIRED — la acción requiere confirmación (ej: 'Cerrar').",
    "ERROR — falla de hardware o modelo.",
]
for e in estados:
    doc.add_paragraph(e, style="List Bullet")

body(doc,
    "Módulos disponibles: Web, Multimedia, Navegación, Accesibilidad. "
    "Confianza mínima configurable (min_confidence, default 0.65). "
    "Si la confianza es baja el sistema lo indica y pide repetir.")
body(doc,
    "El estado se muestra en tiempo real en la pantalla Inicio (sección 'Voz en vivo': "
    "icono de estado, módulo activo, último comando ejecutado, advertencia naranja de "
    "confianza baja) y en el mini overlay al minimizar.")

h2(doc, "Interfaz de usuario (v0.3)")
body(doc,
    "La UI usa CustomTkinter con diseño oscuro (paleta #0D1117, acento #10B981). "
    "Ventana principal con sidebar fija (180 px) y 4 vistas internas:")
vistas = [
    "Inicio: preview de cámara (hilo dedicado), badge de gesto detectado, métricas "
     "faciales (boca, ojos, cejas, sonrisa), estado de módulos (gestos/mouse/voz), "
     "panel 'Voz en vivo' (estado / módulo / comando / advertencia de confianza).",
    "Configuración: tabs Gestos→Acciones / Perfiles / Mouse. Panel derecho con toggle "
     "de Control de voz, resumen de mouse (estado/velocidad/zona muerta/suavizado), "
     "chip de perfil activo y lista de otros perfiles.",
    "Calibración: grid 2-col de tarjetas con slider por gesto, valor en vivo, selector "
     "de perfil. Botones: Guardar / Restaurar defaults / Auto-calibrar / Foto base neutral.",
    "Voz: módulos de activación (Web/Multimedia/Navegación/Accesibilidad con conteo "
     "de comandos), sliders de ganancia y cooldown, diagnóstico de micrófono y modelos.",
]
for v in vistas:
    doc.add_paragraph(v, style="List Bullet")

body(doc,
    "Mini overlay (CTkToplevel al minimizar): preview de cámara en vivo + badge de "
    "gesto detectado + panel de estado de voz (icono de estado + módulo activo). Arrastrable.")

h2(doc, "Configuración y perfiles")
body(doc,
    "La configuración se persiste en config/user_settings.json (merge sobre "
    "default_config.json). Cada perfil contiene: gesture_thresholds, gesture_actions, "
    "voice_actions, mouse_settings. Los cambios se aplican sin reiniciar la app. "
    "Se pueden clonar perfiles.")

h2(doc, "Arquitectura de archivos clave")
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
tbl.rows[0].cells[0].text = "Archivo"
tbl.rows[0].cells[1].text = "Rol"
archivos = [
    ("ui/main_window.py", "Ventana principal (~1900 líneas): UI + bucle de cámara + lógica de gestos/voz"),
    ("ui/settings_store.py", "Capa de acceso a config: get/set de perfiles, umbrales, mouse, voz"),
    ("vision/gesture_detection/facial_gestures.py", "FacialGestureDetector con calibración neutral adaptativa"),
    ("vision/face_tracking/face_detector.py", "MediaPipeFaceDetector (landmarks + blendshapes + nose_tip)"),
    ("core/input/mouse_controller.py", "MouseController: zona muerta, velocidad, suavizado EMA"),
    ("core/input/blink_click_controller.py", "BlinkClickController: guiños unilaterales → clicks"),
    ("core/voice_control/voice_control.py", "VoiceCommandController: máquina de estados de voz"),
    ("core/gesture_engine/hotkey_executor.py", "Ejecución de hotkeys/clicks con pydirectinput"),
    ("config/default_config.json", "Config por defecto: perfiles, umbrales, voz, mouse"),
]
for a, r in archivos:
    row = tbl.add_row()
    row.cells[0].text = a
    row.cells[1].text = r

doc.save("E:/Feria/GazeDash/Cómo Funciona_0.3.docx")
print("Cómo Funciona OK")

# ─────────────────────────────────────────────────────────────────────────────
# DOC 2: Planificación 0.3
# ─────────────────────────────────────────────────────────────────────────────
doc2 = Document()

h1(doc2, "GazeDash — Planificación (v0.3)")
body(doc2, f"Última actualización: {datetime.date.today().strftime('%d/%m/%Y')}")

h2(doc2, "Problema y enfoque vigente")
body(doc2,
    "El problema central es ofrecer una forma accesible y económica de interactuar con "
    "una computadora usando hardware común (cámara web). La solución actual prioriza "
    "gestos faciales configurables y movimiento de nariz como joystick alternativo al mouse.")
body(doc2,
    "El foco técnico ya no es solo seguimiento ocular. El control funcional del cursor "
    "se apoya en la nariz como joystick; la estimación de mirada queda como línea futura.")

h2(doc2, "Alcance real del MVP actual (v0.3)")
tbl2 = doc2.add_table(rows=1, cols=3)
tbl2.style = "Table Grid"
tbl2.rows[0].cells[0].text = "Funcionalidad"
tbl2.rows[0].cells[1].text = "Estado"
tbl2.rows[0].cells[2].text = "Detalle"
features = [
    ("Aplicación principal", "✅ Hecha",
     "Ventana CustomTkinter con sidebar + 4 vistas embebidas (Inicio, Config, Calibración, Voz), mini overlay arrastrable."),
    ("Detección facial", "✅ Hecha",
     "MediaPipe Face Landmarker con landmarks, blendshapes y métricas relevantes."),
    ("Gestos faciales", "✅ Hecha / en ajuste",
     "Detecta boca, sonrisa, cejas, guiño, ojos abiertos y nariz; debounce y calibración neutral adaptativa."),
    ("Mapeo gesto → acción", "✅ Hecha",
     "Acciones por perfil usando hotkeys con pydirectinput."),
    ("Mouse por nariz", "✅ Hecha",
     "Movimiento relativo con centro recalibrable, zona muerta, velocidad y suavizado. Toggle en Config."),
    ("Clic por guiño", "✅ Hecha",
     "Click izquierdo/derecho con guiños unilaterales y cooldown."),
    ("Perfiles", "✅ Hecha",
     "Clonar perfiles, guardar umbrales, acciones, voz y mouse. Cambios inmediatos."),
    ("Calibración de umbrales", "✅ Hecha",
     "Vista embebida con grid de sliders por gesto, guardar/restaurar/auto-calibrar."),
    ("Feedback visual de voz", "✅ Hecho",
     "Panel 'Voz en vivo' en Inicio + panel en mini overlay: estado/módulo/comando/advertencia de confianza baja."),
    ("Control de voz", "⚠️ Parcial",
     "Arquitectura ML + máquina de estados completa. Requiere modelos .pkl en assets/models/modelos_v2/."),
    ("Toggle voz desde Config", "✅ Hecho",
     "CTkSwitch en panel derecho de Configuración, persiste y aplica inmediatamente."),
    ("Toggle mouse desde Config", "✅ Hecho",
     "CTkSwitch en tab Mouse, persiste y aplica inmediatamente."),
    ("Control por mirada real", "🔲 Pendiente",
     "Módulos placeholder. No es parte del MVP actual."),
    ("Pruebas automatizadas", "🔲 Pendiente",
     "No hay suite de tests."),
]
for f, s, d in features:
    row = tbl2.add_row()
    row.cells[0].text = f
    row.cells[1].text = s
    row.cells[2].text = d

h2(doc2, "Historias de usuario")
tbl3 = doc2.add_table(rows=1, cols=4)
tbl3.style = "Table Grid"
tbl3.rows[0].cells[0].text = "ID"
tbl3.rows[0].cells[1].text = "Historia"
tbl3.rows[0].cells[2].text = "Prioridad"
tbl3.rows[0].cells[3].text = "Estado"
hus = [
    ("HU1", "Como usuario quiero ver la cámara y el gesto detectado.", "Alta", "✅ Hecha"),
    ("HU2", "Como usuario quiero controlar el cursor con movimientos de nariz.", "Alta", "✅ Hecha"),
    ("HU3", "Como usuario quiero hacer clic con guiños.", "Alta", "✅ Hecha"),
    ("HU4", "Como usuario quiero asignar gestos faciales a combinaciones de teclas.", "Alta", "✅ Hecha"),
    ("HU5", "Como usuario quiero perfiles distintos para cambiar sensibilidad y acciones.", "Media", "✅ Hecha"),
    ("HU6", "Como usuario quiero recalibrar el centro del mouse al cambiar de postura.", "Alta", "✅ Hecha"),
    ("HU7", "Como usuario quiero ajustar umbrales para evitar activaciones involuntarias.", "Alta", "✅ Hecha"),
    ("HU8", "Como usuario quiero feedback visual del estado del sistema de voz.", "Alta", "✅ Hecha"),
    ("HU9", "Como usuario quiero habilitar/deshabilitar voz y mouse desde Config.", "Alta", "✅ Hecha"),
    ("HU10", "Como usuario quiero comandos de voz opcionales para acciones frecuentes.", "Media", "⚠️ Parcial"),
    ("HU11", "Como usuario quiero controlar el cursor con mirada real.", "Media", "🔲 Pendiente"),
    ("HU12", "Como dev quiero pruebas automatizadas de gestos, config y acciones.", "Alta técnica", "🔲 Pendiente"),
]
for row_data in hus:
    row = tbl3.add_row()
    for i, c in enumerate(row_data):
        row.cells[i].text = c

h2(doc2, "Requerimientos funcionales actualizados")
tbl4 = doc2.add_table(rows=1, cols=4)
tbl4.style = "Table Grid"
tbl4.rows[0].cells[0].text = "ID"
tbl4.rows[0].cells[1].text = "Requerimiento"
tbl4.rows[0].cells[2].text = "Prioridad"
tbl4.rows[0].cells[3].text = "Estado"
rfs = [
    ("RF1", "Abrir app con preview de cámara y estado del controlador.", "Alta", "✅ Cumplido"),
    ("RF2", "Detectar rostro y métricas faciales en tiempo real.", "Alta", "✅ Cumplido"),
    ("RF3", "Detectar gestos faciales sobre una línea base neutral.", "Alta", "✅ Cumplido con ajustes"),
    ("RF4", "Ejecutar hotkeys configurables al detectar gestos.", "Alta", "✅ Cumplido"),
    ("RF5", "Mover cursor por nariz como joystick asistivo.", "Alta", "✅ Cumplido"),
    ("RF6", "Ejecutar click por guiño.", "Alta", "✅ Cumplido"),
    ("RF7", "Guardar perfiles, umbrales, acciones y parámetros del mouse.", "Alta", "✅ Cumplido"),
    ("RF8", "Calibrar umbrales desde la UI.", "Media", "✅ Cumplido — vista embebida"),
    ("RF9", "Mostrar feedback visual del módulo de voz en tiempo real.", "Alta", "✅ Cumplido"),
    ("RF10", "Activar/desactivar voz y mouse desde Configuración.", "Alta", "✅ Cumplido"),
    ("RF11", "Activar comandos de voz y mapearlos a acciones.", "Media", "⚠️ Parcial"),
    ("RF12", "Controlar cursor por mirada estimada.", "Media", "🔲 Pendiente"),
]
for row_data in rfs:
    row = tbl4.add_row()
    for i, c in enumerate(row_data):
        row.cells[i].text = c

h2(doc2, "Backlog priorizado")
tbl5 = doc2.add_table(rows=1, cols=3)
tbl5.style = "Table Grid"
tbl5.rows[0].cells[0].text = "Prioridad"
tbl5.rows[0].cells[1].text = "Tarea"
tbl5.rows[0].cells[2].text = "Motivo"
backlog = [
    ("P0", "Validar carga de modelos de voz en entorno limpio",
     "El módulo de voz requiere .pkl en assets/models/modelos_v2/."),
    ("P0", "Suite mínima de tests (config, mapeo de gestos, hotkeys simuladas)",
     "Reduce regresiones en partes críticas."),
    ("P1", "Hacer configurable el umbral de click por guiño desde perfil",
     "Hoy tiene umbrales fijos en código."),
    ("P1", "Revisar nombres internos de métricas brow_raise/brow_frown",
     "Riesgo de gestos que nunca disparen por nombres inconsistentes."),
    ("P1", "Unificar ruta principal (main_window) y ruta experimental (AppController)",
     "Evita confusión entre UI principal y controlador OpenCV standalone."),
    ("P2", "Implementar o retirar activation_word de voz",
     "Aparece en config pero no tiene efecto hoy."),
    ("P2", "Medición formal de latencia y consumo de CPU",
     "Permite validar requerimientos no funcionales."),
    ("P3", "Estimación de mirada real y calibración de pantalla",
     "Extensión futura, no parte del MVP actual."),
]
for row_data in backlog:
    row = tbl5.add_row()
    for i, c in enumerate(row_data):
        row.cells[i].text = c

h2(doc2, "Criterios de aceptación para la próxima entrega")
criterios = [
    "La app arranca desde python -m app.main en un entorno limpio documentado.",
    "La cámara muestra rostro/malla y gesto detectado sin errores visibles.",
    "Un perfil puede mapear al menos tres gestos a hotkeys y persistirlos.",
    "El mouse por nariz se recalibra y respeta los ajustes del perfil activo.",
    "Los toggles de Mouse y Voz en Configuración aplican inmediatamente.",
    "El panel 'Voz en vivo' muestra estado, módulo y comando en tiempo real.",
    "Los clics por guiño se pueden activar de forma consistente.",
    "README, requirements.txt y estos documentos cuentan la misma verdad del proyecto.",
]
for c in criterios:
    doc2.add_paragraph(c, style="List Bullet")

h2(doc2, "Riesgos y decisiones pendientes")
tbl6 = doc2.add_table(rows=1, cols=3)
tbl6.style = "Table Grid"
tbl6.rows[0].cells[0].text = "Riesgo"
tbl6.rows[0].cells[1].text = "Impacto"
tbl6.rows[0].cells[2].text = "Mitigación recomendada"
riesgos = [
    ("Dependencias de voz faltantes o versiones incorrectas",
     "Módulo de voz queda deshabilitado en entorno nuevo.",
     "Probar instalación limpia y documentar versiones exactas."),
    ("Umbrales calibrados para una sola persona",
     "Falsos positivos en otros usuarios.",
     "Calibración guiada en primera ejecución y perfiles por usuario."),
    ("Acciones pydirectinput disparan sobre el sistema real durante pruebas",
     "Acciones no deseadas.",
     "Modo simulación/test sin ejecución real y confirmaciones para acciones peligrosas."),
    ("Código experimental (AppController) mezclado con ruta principal",
     "Mantenimiento más difícil.",
     "Separar demo/experimental de la app principal."),
]
for row_data in riesgos:
    row = tbl6.add_row()
    for i, c in enumerate(row_data):
        row.cells[i].text = c

doc2.save("E:/Feria/GazeDash/Planificación_0.3.docx")
print("Planificación OK")
